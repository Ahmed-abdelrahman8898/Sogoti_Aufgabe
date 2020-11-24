class test_two:
    def  __init__(self,chrome_path,folder_path):
        self.chrome=chrome_path
        self.path=folder_path
    

        from selenium import webdriver
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.common.exceptions import TimeoutException
        from selenium.webdriver.common.keys import Keys
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import date
        from docx.shared import RGBColor
        import pandas as pd
        import sys
        import random
        import string
        try:
            self.out_passed=True
            testdata_path=self.path +'\\' +'2_testcase.xlsx'
            Testdata=pd.read_excel(testdata_path)
            Step=0
            today = date.today()
            d4 = today.strftime("%b-%d-%Y")
            d3=today.strftime("%b_%d_%Y")
            document = Document()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title = p.add_run("2_testcase")
            title.bold = True
            title.font.size = Pt(16)
            Para1=document.add_paragraph()
            title1 = Para1.add_run(d4)
            title1.bold = True
            title1.font.size = Pt(12)
            title1.font.color.rgb = RGBColor(128,0,0)
            PathChrome=self.chrome
            driver=webdriver.Chrome (PathChrome)
            driver.get(" https://www.sogeti.com/")
            driver.maximize_window()
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'//*[@id="header"]/div[1]/nav[1]/ul[1]/li[3]/div[1]/span[1]')))
            document.add_paragraph(Testdata.at[Step,'Describtion'])
            document.add_paragraph('Passed')
            driver.save_screenshot(self.path +'\\' +'pic.jpg')
            document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
            Step=Step+1
            driver.find_element_by_xpath('//*[@id="header"]/div[1]/nav[1]/ul[1]/li[3]/div[1]/span[1]').click()
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'//*[@id="header"]/div[1]/div[5]/ul[1]/li[7]/a[1]')))
            driver.find_element_by_xpath('//*[@id="header"]/div[1]/div[5]/ul[1]/li[7]/a[1]').click()
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'//*[@id="primary_content"]/div[1]/div[2]/div[1]/h1[1]/span[1]')))
            r_name1=''
            r_name2=''
            r_number=''
            for x in range(6):
                r_name1=r_name1+random.choice(string.ascii_letters)
            driver.find_element_by_xpath('//*[@id="4ff2ed4d-4861-4914-86eb-87dfa65876d8"]').send_keys(r_name1)
            for x in range(6):
                r_name2=r_name2+random.choice(string.ascii_letters)
            driver.find_element_by_xpath('//*[@id="11ce8b49-5298-491a-aebe-d0900d6f49a7"]').send_keys(r_name2)
            
            driver.find_element_by_xpath('//*[@id="056d8435-4d06-44f3-896a-d7b0bf4d37b2"]').send_keys(r_name1+'.'+r_name2+'@rrr.com')
            for x in range(8):
                r_number=r_number+str(random.randint(0, 9))
            driver.find_element_by_xpath('//*[@id="755aa064-7be2-432b-b8a2-805b5f4f9384"]').send_keys(r_number)
            driver.find_element_by_xpath('//*[@id="88459d00-b812-459a-99e4-5dc6eff2aa19"]').send_keys('Hello '+r_name1+' '+' '+r_name2)
            driver.find_element_by_xpath('//*[@value="I agree"]').click()
            document.add_paragraph(Testdata.at[Step,'Describtion'])
            document.add_paragraph('Passed')
            driver.save_screenshot(self.path +'\\' +'pic.jpg')
            document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
            Step=Step+1
            driver.find_element_by_xpath('//*[@name="submit"]').click()
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'//*[@class="Form__Status__Message Form__Success__Message"]')))
            message=driver.find_element_by_xpath('//*[@class="Form__Status__Message Form__Success__Message"]').text
            if(message=='Thank you for contacting us.'):
                document.add_paragraph(Testdata.at[Step,'Describtion'])
                document.add_paragraph('Passed')
                driver.save_screenshot(self.path +'\\' +'pic.jpg')
                document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                Step=Step+1
            else:
                document.add_paragraph(Testdata.at[Step,'Describtion'])
                document.add_paragraph('Failed:Wrong Titel')
                driver.save_screenshot(self.path +'\\' +'pic.jpg')
                document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                raise Exception("BusinessError:wronge message")
                
        except:
            document.add_paragraph('sys.exc_info()')
            driver.save_screenshot(self.path +'\\' +'pic.jpg')
            self.out_passed=False
        finally:
            docName=self.path +'\\'+'2_testcase'+'_'+d3+'.docx'
            document.save(docName)
            self.report_path=docName
            self.report=[self.out_passed,self.report_path]
            driver.quit()
    def result (self):
        
        return self.report
