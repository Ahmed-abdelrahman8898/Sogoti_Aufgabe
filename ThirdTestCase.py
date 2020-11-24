class test_three:
    def  __init__(self,chrome_path,folder_path):
        self.chrome=chrome_path
        self.path=folder_path
        from selenium import webdriver
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.common.exceptions import TimeoutException
        from selenium.webdriver.common.keys import Keys
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import date
        from docx.shared import RGBColor
        import pandas as pd
        import sys
        try:
            self.out_passed=True
            Testdata=pd.read_excel(self.path +'\\' +'3_testcase.xlsx')
            Step=0
            today = date.today()
            d4 = today.strftime("%b-%d-%Y")
            d3=today.strftime("%b_%d_%Y")
            document = Document()
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title = p.add_run("3_testcase")
            title.bold = True
            title.font.size = Pt(16)
            Para1=document.add_paragraph()
            title1 = Para1.add_run(d4)
            title1.bold = True
            title1.font.size = Pt(12)
            title1.font.color.rgb = RGBColor(128,0,0)
            PathChrome=self.chrome
            driver=webdriver.Chrome (PathChrome)
            driver.get(" https://www.sogeti.com/")
            driver.maximize_window()
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'//*[@id="header"]/div[1]/nav[1]/ul[1]/li[3]/div[1]/span[1]')))
            document.add_paragraph(Testdata.at[Step,'Describtion'])
            document.add_paragraph('Passed')
            driver.save_screenshot(self.path +'\\' +'pic.jpg')
            document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
            Step=Step+1
            WebDriverWait(driver, 20).until(EC.visibility_of_element_located((By.XPATH,'//*[@id="header"]/div[1]/nav[1]/ul[1]/li[3]/div[1]/span[1]')))
            driver.find_element_by_xpath('//*[@class="sprite-header sprite-globe"]').click()
            for x in range (12):
                x=x+1
                MainWin=driver.window_handles[0]
                driver.switch_to.window(MainWin)
                name=driver.find_element_by_xpath('//*[@class="country-list"]/ul[1]/li[%d]'%(x)).text
                driver.find_element_by_xpath('//*[@class="country-list"]/ul[1]/li[%d]'%(x)).click()
                w=driver.window_handles[1]
                driver.switch_to.window(w)
                
                if (name.lower() in driver.title.lower() ):
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Passed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                elif('Deutschland' in driver.title):
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Passed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                elif('.nl' in driver.current_url ):
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Passed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                elif('.no' in driver.current_url ):
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Passed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                elif('.es' in driver.current_url ):
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Passed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                elif('.se' in driver.current_url ):
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Passed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                else:
                    document.add_paragraph(Testdata.at[Step,'Describtion'])
                    document.add_paragraph('Failed - '+name)
                    driver.save_screenshot(self.path +'\\' +'pic.jpg')
                    document.add_picture(self.path +'\\' +'pic.jpg',width=Inches(7))
                    self.out_passed=False
                driver.close()
        except:
            document.add_paragraph('sys.exc_info()')
            driver.save_screenshot('pic.jpg')
            self.out_passed=False
            
        finally:
            docName=self.path +'\\'+'3_testcase'+'_'+d3+'.docx'
            document.save(docName)
            self.report_path=docName
            self.report=[self.out_passed,self.report_path]
            driver.quit()
    def result (self):
        
        return self.report