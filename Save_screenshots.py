from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.keys import Keys
driver=webdriver.Chrome (r"C:\Users\ahmedhussienalwany.a\Downloads\chromedriver_win32\chromedriver.exe")
driver.get("https://www.menalite.com/menaitech/application/hrms/index.php")
driver.maximize_window()
driver.save_screenshot('nos.jpg')
username = driver.find_element_by_xpath('//*[@id="userid"]')
username.screenshot('nos1.jpg')
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.shared import RGBColor
today = date.today()
d4 = today.strftime("%b-%d-%Y")
document = Document()
p = document.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title = p.add_run("Test case 1")
title.bold = True
title.font.size = Pt(16)
d=document.add_paragraph()
title1 = d.add_run(d4)
title1.bold = True
title1.font.size = Pt(12)
title1.font.color.rgb = RGBColor(128,0,0)

document.add_picture('nos.jpg',width=Inches(7))
#p = document.add_paragraph()
#p.add_run("titel")
document.add_paragraph("yala bena")
document.add_picture('nos.jpg',width=Inches(7))
document.save('my_report.docx')
